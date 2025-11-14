import os
import logging
from typing import Optional

# Import the base classes from the document processor
from ..pii_anonymizer.document_processor import DocumentHandler, PIIManager

# Set up logging
logger = logging.getLogger('pii_anonymizer.handlers')


class DocxDocumentHandler(DocumentHandler):
    """Handler for Microsoft Word (.docx) documents."""

    def __init__(self, pii_manager: PIIManager):
        """Initialize the DOCX document handler."""
        super().__init__(pii_manager)
        self._check_dependencies()

    def _check_dependencies(self) -> None:
        """Check if required dependencies are installed."""
        try:
            import docx
            logger.info("python-docx is installed.")
        except ImportError:
            logger.warning("python-docx is not installed. Please install it with: pip install python-docx")
            raise ImportError("python-docx is required for DOCX processing.")

    def can_handle(self, file_path: str) -> bool:
        """Check if this handler can process the given file."""
        _, ext = os.path.splitext(file_path.lower())
        return ext == '.docx'

    def extract_text(self, file_path: str) -> str:
        """Extract text from a DOCX file."""
        try:
            from docx import Document

            doc = Document(file_path)
            full_text = []

            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                full_text.append(paragraph.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        full_text.append(cell.text)

            extracted_text = '\n'.join(full_text)
            logger.info(f"Extracted text from DOCX: {file_path}")
            return extracted_text

        except ImportError:
            logger.error("python-docx is not installed. Cannot extract text from DOCX.")
            raise ImportError("python-docx is required for DOCX processing.")
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {str(e)}")
            raise

    def anonymize(self, input_path: str, output_path: str) -> None:
        """Anonymize a DOCX document."""
        # Extract text
        text = self.extract_text(input_path)

        # Anonymize text
        anonymized_text = self.pii_manager.anonymize_text(text)

        # Save as plain text file (with .txt extension)
        if not output_path.endswith('.txt'):
            output_path = output_path.rsplit('.', 1)[0] + '.txt'

        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(anonymized_text)

        logger.info(f"Anonymized DOCX document: {input_path} -> {output_path}")

        # Log statistics
        stats = self.pii_manager.get_statistics()
        logger.info(f"PII Statistics: {stats}")

    def restore(self, anonymized_path: str, output_path: str) -> None:
        """Restore an anonymized DOCX document."""
        # Read anonymized text (should be .txt file)
        try:
            with open(anonymized_path, 'r', encoding='utf-8') as f:
                anonymized_text = f.read()
        except UnicodeDecodeError:
            with open(anonymized_path, 'r', encoding='latin-1') as f:
                anonymized_text = f.read()

        # Restore text
        restored_text = self.pii_manager.restore_text(anonymized_text)

        # Save as plain text file
        if not output_path.endswith('.txt'):
            output_path = output_path.rsplit('.', 1)[0] + '.txt'

        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(restored_text)

        logger.info(f"Restored DOCX document: {anonymized_path} -> {output_path}")


class XlsxDocumentHandler(DocumentHandler):
    """Handler for Microsoft Excel (.xlsx) spreadsheets."""

    def __init__(self, pii_manager: PIIManager):
        """Initialize the XLSX document handler."""
        super().__init__(pii_manager)
        self._check_dependencies()

    def _check_dependencies(self) -> None:
        """Check if required dependencies are installed."""
        try:
            import openpyxl
            logger.info("openpyxl is installed.")
        except ImportError:
            logger.warning("openpyxl is not installed. Please install it with: pip install openpyxl")
            raise ImportError("openpyxl is required for XLSX processing.")

    def can_handle(self, file_path: str) -> bool:
        """Check if this handler can process the given file."""
        _, ext = os.path.splitext(file_path.lower())
        return ext == '.xlsx'

    def extract_text(self, file_path: str) -> str:
        """Extract text from an XLSX file."""
        try:
            from openpyxl import load_workbook

            workbook = load_workbook(filename=file_path, read_only=True, data_only=True)
            full_text = []

            # Process each sheet
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                full_text.append(f"\n=== Sheet: {sheet_name} ===\n")

                # Extract all cell values
                for row in sheet.iter_rows(values_only=True):
                    row_values = [str(cell) if cell is not None else '' for cell in row]
                    # Only add non-empty rows
                    if any(val.strip() for val in row_values):
                        full_text.append('\t'.join(row_values))

            workbook.close()
            extracted_text = '\n'.join(full_text)
            logger.info(f"Extracted text from XLSX: {file_path}")
            return extracted_text

        except ImportError:
            logger.error("openpyxl is not installed. Cannot extract text from XLSX.")
            raise ImportError("openpyxl is required for XLSX processing.")
        except Exception as e:
            logger.error(f"Error extracting text from XLSX: {str(e)}")
            raise

    def anonymize(self, input_path: str, output_path: str) -> None:
        """Anonymize an XLSX spreadsheet."""
        # Extract text
        text = self.extract_text(input_path)

        # Anonymize text
        anonymized_text = self.pii_manager.anonymize_text(text)

        # Save as plain text file (with .txt extension)
        if not output_path.endswith('.txt'):
            output_path = output_path.rsplit('.', 1)[0] + '.txt'

        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(anonymized_text)

        logger.info(f"Anonymized XLSX spreadsheet: {input_path} -> {output_path}")

        # Log statistics
        stats = self.pii_manager.get_statistics()
        logger.info(f"PII Statistics: {stats}")

    def restore(self, anonymized_path: str, output_path: str) -> None:
        """Restore an anonymized XLSX spreadsheet."""
        # Read anonymized text (should be .txt file)
        try:
            with open(anonymized_path, 'r', encoding='utf-8') as f:
                anonymized_text = f.read()
        except UnicodeDecodeError:
            with open(anonymized_path, 'r', encoding='latin-1') as f:
                anonymized_text = f.read()

        # Restore text
        restored_text = self.pii_manager.restore_text(anonymized_text)

        # Save as plain text file
        if not output_path.endswith('.txt'):
            output_path = output_path.rsplit('.', 1)[0] + '.txt'

        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(restored_text)

        logger.info(f"Restored XLSX spreadsheet: {anonymized_path} -> {output_path}")
